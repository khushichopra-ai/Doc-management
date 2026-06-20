from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET


@dataclass(slots=True)
class ParsedDocument:
    text: str
    extension: str


class DocumentParser:
    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        extension = Path(filename).suffix.lower().lstrip(".")
        if extension == "txt":
            return ParsedDocument(text=content.decode("utf-8", errors="ignore"), extension=extension)
        if extension == "docx":
            return ParsedDocument(text=self._parse_docx(content), extension=extension)
        if extension == "xlsx":
            return ParsedDocument(text=self._parse_xlsx(content), extension=extension)
        if extension == "pptx":
            return ParsedDocument(text=self._parse_pptx(content), extension=extension)
        if extension == "pdf":
            return ParsedDocument(text=self._parse_pdf(content), extension=extension)
        raise ValueError("Unsupported file type.")

    def _parse_docx(self, content: bytes) -> str:
        # Mandated library: python-docx.
        import docx  # type: ignore

        document = docx.Document(io.BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    parts.append(" ".join(cells))
        return self._clean_text("\n".join(parts))

    def _parse_xlsx(self, content: bytes) -> str:
        # Mandated library: openpyxl.
        from openpyxl import load_workbook  # type: ignore

        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        rows: list[str] = []
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    rows.append(" ".join(values))
        workbook.close()
        return self._clean_text("\n".join(rows))

    def _parse_pptx(self, content: bytes) -> str:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            slide_names = [name for name in archive.namelist() if name.startswith("ppt/slides/slide")]
            texts: list[str] = []
            for slide_name in slide_names:
                root = ET.fromstring(archive.read(slide_name))
                namespace = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
                for text_node in root.findall(".//a:t", namespace):
                    if text_node.text:
                        texts.append(text_node.text)
        return self._clean_text("\n".join(texts))

    def _parse_pdf(self, content: bytes) -> str:
        try:
            import fitz  # type: ignore
        except Exception as exc:
            raise ValueError("PDF parsing requires PyMuPDF (fitz) to be installed.") from exc

        text_parts: list[str] = []
        with fitz.open(stream=content, filetype="pdf") as document:
            for page in document:
                text_parts.append(page.get_text("text"))
        return self._clean_text("\n".join(text_parts))

    def _clean_text(self, text: str) -> str:
        lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.isdigit():
                continue
            lines.append(stripped)
        return "\n".join(lines)

